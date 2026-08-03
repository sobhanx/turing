"""DOCX exporter (python-docx) — premium meeting report layout."""

from __future__ import annotations

from typing import BinaryIO

from turing.services.export.base import BaseExporter, ExportRegistry
from turing.services.export.document import ExportDocument
from turing.services.export.font_assets import unicode_font_paths
from turing.services.export import labels as L
from turing.services.export.layout import (
    BRAND_ACCENT,
    BRAND_DECISION_BG,
    BRAND_MUTED,
    BRAND_PRIMARY,
    BRAND_SUMMARY_BG,
    BRAND_SURFACE,
    speaker_color,
)


@ExportRegistry.register
class DOCXExporter(BaseExporter):
    format_code = "docx"
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    file_extension = "docx"
    label = "Word (DOCX)"

    def write(self, document: ExportDocument, output: BinaryIO) -> None:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError(
                "DOCX export requires python-docx. Install django-turing[export]."
            ) from exc

        unicode_font_paths()
        font_name = "DejaVu Sans"
        doc = Document()
        rtl = document.rtl
        vis = document.visibility
        align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        center = WD_ALIGN_PARAGRAPH.CENTER
        gen_ts = document.generated_display

        # Narrow professional margins
        for section in doc.sections:
            section.top_margin = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)
            _add_footer(section, gen_ts, rtl=rtl, font_name=font_name)

        def add_runs(paragraph, text: str, *, bold=False, size=11, color=None):
            run = paragraph.add_run(text or "")
            run.bold = bold
            run.font.name = font_name
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if color is not None:
                run.font.color.rgb = RGBColor(*color)
            return run

        def para(text="", *, style=None, align_=None, before=0, after=6, bold=False, size=11, color=None):
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            _set_paragraph_rtl(p, rtl=rtl)
            p.alignment = align_ if align_ is not None else align
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after = Pt(after)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            if text:
                add_runs(p, text, bold=bold, size=size, color=color)
            return p

        def heading(text: str, level: int = 1):
            h = doc.add_heading(text, level=level)
            _set_paragraph_rtl(h, rtl=rtl)
            h.alignment = align
            for run in h.runs:
                run.font.name = font_name
                run.font.color.rgb = RGBColor(*_hex_rgb(BRAND_PRIMARY))
            return h

        def add_kv_table(rows: list[tuple[str, str]]):
            if not rows:
                return
            table = doc.add_table(rows=len(rows), cols=2)
            table.style = "Table Grid"
            for i, (label, value) in enumerate(rows):
                cell_l, cell_v = table.rows[i].cells
                _fill_cell(cell_l, label, font_name, muted=True, rtl=rtl, shade=BRAND_SURFACE)
                _fill_cell(cell_v, value or "—", font_name, rtl=rtl)
            para("", after=8)

        # ----- Cover -----
        para(
            L.BRAND_NAME,
            align_=center,
            after=4,
            size=9,
            color=_hex_rgb(BRAND_MUTED),
        )
        title = doc.add_heading(L.REPORT_TITLE, level=0)
        _set_paragraph_rtl(title, rtl=rtl)
        title.alignment = center
        for run in title.runs:
            run.font.name = font_name
            run.font.color.rgb = RGBColor(*_hex_rgb(BRAND_PRIMARY))

        para(
            document.cover_title(),
            align_=center,
            after=12,
            size=12,
            color=_hex_rgb(BRAND_ACCENT),
        )
        logo = para(L.LOGO_PLACEHOLDER, align_=center, after=14, size=9, color=_hex_rgb(BRAND_MUTED))
        _shade_paragraph(logo, BRAND_SURFACE)

        add_kv_table(document.cover_rows())

        # ----- Meeting Information -----
        meeting_rows = document.meeting_info_rows()
        if meeting_rows:
            heading(L.SECTION_MEETING_INFO, level=1)
            add_kv_table(meeting_rows)

        # ----- Executive Summary -----
        if vis.show_ai_summary:
            heading(L.SECTION_EXECUTIVE_SUMMARY, level=1)
            summary_text = (document.summary or "").strip() or L.EMPTY_SUMMARY
            sp = para(summary_text, after=10, size=11)
            _shade_paragraph(sp, BRAND_SUMMARY_BG)

        # ----- Key Topics -----
        if vis.show_key_topics:
            heading(L.SECTION_KEY_TOPICS, level=1)
            if document.topics:
                for topic in document.topics:
                    p = doc.add_paragraph(style="List Bullet")
                    _set_paragraph_rtl(p, rtl=rtl)
                    p.alignment = align
                    add_runs(p, topic, size=11)
            else:
                para(L.EMPTY_TOPICS, size=10, color=_hex_rgb(BRAND_MUTED))

        # ----- Action Items -----
        if vis.show_action_items:
            heading(L.SECTION_ACTION_ITEMS, level=1)
            if document.action_items:
                for item in document.action_items:
                    line = f"☐  {item.task}"
                    if item.owner:
                        line += f"  —  {item.owner}"
                    if item.deadline:
                        line += f"  ({item.deadline})"
                    para(line, after=4, size=11)
            else:
                para(L.EMPTY_ACTION_ITEMS, size=10, color=_hex_rgb(BRAND_MUTED))

        # ----- Decisions -----
        if vis.show_decisions:
            heading(L.SECTION_DECISIONS, level=1)
            if document.decisions:
                for d in document.decisions:
                    dp = para(f"•  {d}", after=4, size=11)
                    _shade_paragraph(dp, BRAND_DECISION_BG)
            else:
                para(L.EMPTY_DECISIONS, size=10, color=_hex_rgb(BRAND_MUTED))

        # ----- Keywords -----
        if vis.show_keywords:
            heading(L.SECTION_KEYWORDS, level=1)
            if document.keywords:
                kp = para("", after=10)
                for i, kw in enumerate(document.keywords):
                    if i:
                        add_runs(kp, "   ", size=9)
                    add_runs(
                        kp,
                        f"  {kw}  ",
                        size=9,
                        color=_hex_rgb(BRAND_ACCENT),
                    )
            else:
                para(L.EMPTY_KEYWORDS, size=10, color=_hex_rgb(BRAND_MUTED))

        # ----- Transcript -----
        if vis.show_full_transcript:
            heading(L.SECTION_TRANSCRIPT, level=1)
            speaker_index: dict[str, int] = {}
            next_idx = 0
            if document.turns:
                for turn in document.turns:
                    name = turn.speaker_name or "—"
                    if name not in speaker_index:
                        speaker_index[name] = next_idx
                        next_idx += 1
                    color = _hex_rgb(speaker_color(speaker_index[name]))
                    # Speaker label above exact segment text (UI order).
                    para(name, before=12, after=2, bold=True, size=11, color=color)
                    ts = document.turn_timestamp(turn.start_display)
                    if ts:
                        para(
                            ts,
                            after=2,
                            size=8,
                            color=_hex_rgb(BRAND_MUTED),
                        )
                    # Preserve exact segment text, including blank lines.
                    lines = (turn.text or "").split("\n")
                    if not lines:
                        lines = [""]
                    for i, line in enumerate(lines):
                        is_last = i == len(lines) - 1
                        tp = para(
                            line,
                            after=10 if is_last else 0,
                            size=11,
                        )
                        tp.paragraph_format.line_spacing = 1.15
            else:
                para(L.EMPTY_TRANSCRIPT, size=10, color=_hex_rgb(BRAND_MUTED))

        doc.save(output)


def _hex_rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)


def _set_paragraph_rtl(paragraph, *, rtl: bool) -> None:
    if not rtl:
        return
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    p_pr.append(shd)


def _fill_cell(cell, text: str, font_name: str, *, muted=False, rtl=False, shade=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_rtl(p, rtl=rtl)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text or "—")
    run.font.name = font_name
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if muted:
        run.font.color.rgb = RGBColor(*_hex_rgb(BRAND_MUTED))
    else:
        run.font.color.rgb = RGBColor(*_hex_rgb(BRAND_PRIMARY))
    if shade:
        from docx.oxml import OxmlElement

        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shade.lstrip("#"))
        tc_pr.append(shd)


def _add_footer(section, gen_ts: str, *, rtl: bool, font_name: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rtl(p, rtl=rtl)

    run_l = p.add_run(L.FOOTER_GENERATED_BY)
    run_l.font.name = font_name
    run_l.font.size = Pt(8)
    run_l.font.color.rgb = RGBColor(*_hex_rgb(BRAND_MUTED))

    p.add_run("  ·  ")

    # PAGE field
    run_page = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_page._r.append(fld_char_begin)
    run_page._r.append(instr)
    run_page._r.append(fld_char_end)
    run_page.font.name = font_name
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = RGBColor(*_hex_rgb(BRAND_MUTED))

    p.add_run("  ·  ")
    run_r = p.add_run(gen_ts)
    run_r.font.name = font_name
    run_r.font.size = Pt(8)
    run_r.font.color.rgb = RGBColor(*_hex_rgb(BRAND_MUTED))
